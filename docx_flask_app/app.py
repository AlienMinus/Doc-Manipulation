import os
import re
import base64
import zipfile
import requests
import tempfile
from io import BytesIO
from flask import Flask, render_template, request, send_file, flash, jsonify
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn

app = Flask(__name__)
app.secret_key = 'super_secret_key_for_demo_purposes'

def replace_text_in_doc(doc, search_text, replace_text):
    """
    Simple search and replace for Word documents.
    
    NOTE: This method iterates through paragraphs. If the search text 
    is split across multiple 'runs' (e.g., half bold, half normal), 
    this simple logic might miss it. It is optimized for standard text.
    """
    # 1. Replace in Paragraphs
    for paragraph in doc.paragraphs:
        if search_text in paragraph.text:
            # We iterate runs to attempt to preserve formatting
            for run in paragraph.runs:
                if search_text in run.text:
                    run.text = run.text.replace(search_text, replace_text)
    
    # 2. Replace in Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if search_text in paragraph.text:
                        for run in paragraph.runs:
                            if search_text in run.text:
                                run.text = run.text.replace(search_text, replace_text)

def extract_metadata(doc):
    core_props = doc.core_properties
    return {
        "Author": core_props.author,
        "Created": core_props.created,
        "Last Modified By": core_props.last_modified_by,
        "Title": core_props.title,
        "Revision": core_props.revision
    }

def extract_images_base64(source_stream):
    images = []
    with zipfile.ZipFile(source_stream, 'r') as docx_zip:
        for file in docx_zip.namelist():
            if file.startswith('word/media/') and len(file) > len('word/media/'):
                image_data = docx_zip.read(file)
                encoded_img = base64.b64encode(image_data).decode('utf-8')
                filename = os.path.basename(file)
                # Simple extension check for mime type
                ext = filename.split('.')[-1].lower()
                mime_type = 'image/jpeg' if ext in ['jpg', 'jpeg'] else f'image/{ext}'
                
                images.append({'filename': filename, 'data': encoded_img, 'mime': mime_type})
    return images

def create_doc_from_markdown(md_text, image_map=None):
    if image_map is None:
        image_map = {}
    doc = Document()
    # Regex for markdown image: ![alt](url)
    img_pattern = re.compile(r'^!\[(.*?)\]\((.*?)\)$')

    for line in md_text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        img_match = img_pattern.match(line)
        
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif img_match:
            image_ref = img_match.group(2)
            try:
                if image_ref in image_map:
                    # Use uploaded image
                    image_map[image_ref].seek(0)
                    doc.add_picture(image_map[image_ref], width=Inches(6))
                else:
                    response = requests.get(image_ref)
                    response.raise_for_status()
                    doc.add_picture(BytesIO(response.content), width=Inches(6))
            except Exception:
                doc.add_paragraph(f"[Could not load image: {image_ref}]")
        else:
            doc.add_paragraph(line)
    return doc

def convert_docx_to_markdown(doc):
    lines = []
    # Process Paragraphs
    for p in doc.paragraphs:
        text = p.text.strip()
        
        # Check for images in the paragraph
        blips = p._element.findall('.//' + qn('a:blip'))
        
        if not text and not blips:
            continue
        
        if text:
            style_name = p.style.name
            if style_name.startswith('Heading'):
                try:
                    level = int(style_name.split()[-1])
                    lines.append(f"{'#' * level} {text}")
                except:
                    lines.append(f"**{text}**")
            elif 'List Bullet' in style_name:
                lines.append(f"- {text}")
            elif 'List Number' in style_name:
                lines.append(f"1. {text}")
            else:
                lines.append(text)
        
        # Process Images
        for blip in blips:
            embed_attr = blip.get(qn('r:embed'))
            if embed_attr:
                try:
                    image_part = doc.part.rels[embed_attr].target_part
                    if hasattr(image_part, 'blob'):
                        image_data = image_part.blob
                        content_type = image_part.content_type
                        encoded_img = base64.b64encode(image_data).decode('utf-8')
                        lines.append(f"![Image](data:{content_type};base64,{encoded_img})")
                except (KeyError, ValueError, AttributeError):
                    pass

        lines.append("")
        
    # Process Tables (Append at end)
    if doc.tables:
        lines.append("\n--- Tables ---\n")
        for table in doc.tables:
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    lines.append("| " + " | ".join(['---'] * len(cells)) + " |")
            lines.append("")
            
    return "\n".join(lines)

# --- API Endpoints ---

@app.route('/api/replace', methods=['POST'])
def api_replace():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    search_text = request.form.get('search_text')
    replace_text = request.form.get('replace_text')

    if not search_text or not replace_text:
        return jsonify({'error': 'Missing search_text or replace_text'}), 400

    if file and file.filename.endswith('.docx'):
        try:
            source_stream = BytesIO(file.read())
            doc = Document(source_stream)
            replace_text_in_doc(doc, search_text, replace_text)
            
            target_stream = BytesIO()
            doc.save(target_stream)
            target_stream.seek(0)
            
            return send_file(
                target_stream,
                as_attachment=True,
                download_name=f"modified_{file.filename}",
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    return jsonify({'error': 'Invalid file format. Please upload a .docx file.'}), 400

@app.route('/api/metadata', methods=['POST'])
def api_metadata():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file and file.filename.endswith('.docx'):
        try:
            source_stream = BytesIO(file.read())
            doc = Document(source_stream)
            metadata = extract_metadata(doc)
            return jsonify({'metadata': metadata})
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    return jsonify({'error': 'Invalid file format'}), 400

@app.route('/api/text', methods=['POST'])
def api_text():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file and file.filename.endswith('.docx'):
        try:
            source_stream = BytesIO(file.read())
            doc = Document(source_stream)
            full_text = [p.text for p in doc.paragraphs if p.text.strip()]
            return jsonify({'text': "\n".join(full_text)})
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    return jsonify({'error': 'Invalid file format'}), 400

@app.route('/api/images', methods=['POST'])
def api_images():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file and file.filename.endswith('.docx'):
        try:
            source_stream = BytesIO(file.read())
            images = extract_images_base64(source_stream)
            return jsonify({'images': images})
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    return jsonify({'error': 'Invalid file format'}), 400

@app.route('/api/tables', methods=['POST'])
def api_tables():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file and file.filename.endswith('.docx'):
        try:
            source_stream = BytesIO(file.read())
            doc = Document(source_stream)
            tables_data = []
            for table in doc.tables:
                t_rows = [[cell.text for cell in row.cells] for row in table.rows]
                tables_data.append(t_rows)
            return jsonify({'tables': tables_data})
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    return jsonify({'error': 'Invalid file format'}), 400

@app.route('/api/generate', methods=['POST'])
def api_generate():
    data = request.get_json(silent=True) or request.form
    md_text = data.get('markdown_text', '')
    
    if not md_text:
        return jsonify({'error': 'No markdown_text provided'}), 400
        
    image_map = {}
    if 'images' in request.files:
        for file in request.files.getlist('images'):
            if file.filename:
                image_map[file.filename] = BytesIO(file.read())

    try:
        doc = create_doc_from_markdown(md_text, image_map)
        target_stream = BytesIO()
        doc.save(target_stream)
        target_stream.seek(0)
        return send_file(
            target_stream,
            as_attachment=True,
            download_name="generated_doc.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/docx-to-md', methods=['POST'])
def api_docx_to_md():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file and file.filename.endswith('.docx'):
        try:
            source_stream = BytesIO(file.read())
            doc = Document(source_stream)
            md_text = convert_docx_to_markdown(doc)
            
            if request.form.get('preview') == 'true':
                return jsonify({'markdown': md_text})
            
            target_stream = BytesIO()
            target_stream.write(md_text.encode('utf-8'))
            target_stream.seek(0)
            
            return send_file(
                target_stream,
                as_attachment=True,
                download_name=f"{os.path.splitext(file.filename)[0]}.md",
                mimetype='text/markdown'
            )
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    return jsonify({'error': 'Invalid file format'}), 400

@app.route('/api/pdf-to-docx', methods=['POST'])
def api_pdf_to_docx():
    return jsonify({'error': 'PDF conversion is disabled in this deployment due to serverless size limits.'}), 400

@app.route('/api/docx-to-pdf', methods=['POST'])
def api_docx_to_pdf():
    return jsonify({'error': 'PDF conversion is disabled in this deployment due to serverless size limits.'}), 400

@app.route('/', methods=['GET', 'POST'])
def index():
    result = {}
    form_data = {}
    active_feature = 'text'
    
    if request.method == 'POST':
        form_data = request.form
        feature = form_data.get('feature')
        active_feature = feature

        # Map features to their corresponding API functions
        api_functions = {
            'replace': api_replace,
            'metadata': api_metadata,
            'text': api_text,
            'images': api_images,
            'tables': api_tables,
            'generate': api_generate,
            'docx-to-md': api_docx_to_md,
            'pdf-to-docx': api_pdf_to_docx,
            'docx-to-pdf': api_docx_to_pdf
        }

        if feature in api_functions:
            # Call the API function internally (uses the same global 'request' object)
            response = api_functions[feature]()

            # Handle File Download (Success for Replace/Generate)
            if response.mimetype == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return response
            
            # Handle JSON Response (Data or Error)
            if response.is_json:
                data = response.get_json()
                if response.status_code >= 400:
                    flash(data.get('error', 'An error occurred'))
                else:
                    # Merge API data (e.g., {'text': ...}) into result for template
                    result.update(data)
            else:
                # Fallback for unexpected responses
                if response.status_code >= 400:
                    flash("An error occurred processing the request.")

    return render_template('index.html', result=result, form_data=form_data, active_feature=active_feature)

if __name__ == '__main__':
    # Use waitress for production serving (requires: pip install waitress)
    from waitress import serve
    print("Production server running on http://127.0.0.1:5000")
    serve(app, host='0.0.0.0', port=5000)
